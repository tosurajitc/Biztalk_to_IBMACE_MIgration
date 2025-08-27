#!/usr/bin/env python3
"""
Functional Document Generator for BizTalk to ACE Migration
Smart, small and efficient - generates Word documents
"""

import os
import json
import xml.etree.ElementTree as ET
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any

try:
    from groq import Groq
    from prompt_module import get_functional_documentation_prompt
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError as e:
    print(f"Missing dependency: {e}")
    DOCX_AVAILABLE = False

class FunctionalDocumentGenerator:
    def __init__(self):
        api_key = os.getenv('GROQ_API_KEY')
        self.llm = Groq(api_key=api_key) if api_key else None
        self.config = {"groq_settings": {"model": "llama3-70b-8192", "temperature": 0.1}}
        self.analysis_results = {"processed_files": [], "business_processes": [], 
                               "data_transformations": [], "integration_points": []}
        self.scan_depth = 4
    
    def generate_functional_document(self, biztalk_dir: str, confluence_path: str = None, 
                                   output_dir: str = None) -> str:
        """Main generation method - generates Word document"""
        print("📋 Starting Functional Document Generation...")
        
        if not output_dir:
            output_dir = os.path.join(biztalk_dir, "functional_documentation")
        os.makedirs(output_dir, exist_ok=True)
        
        # Analyze BizTalk files
        biztalk_analysis = self._scan_biztalk_files(biztalk_dir)
        
        # Process Confluence if provided
        confluence_analysis = {}
        if confluence_path and os.path.exists(confluence_path):
            confluence_analysis = self._process_confluence(confluence_path)
        
        # Generate document content - 100% LLM ONLY
        document_content = self._generate_with_llm(biztalk_analysis, confluence_analysis)
        
        # Save as Word document with folder-specific name
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        folder_name = Path(biztalk_dir).name  # Get folder name for document naming
        
        if DOCX_AVAILABLE:
            doc_path = self._create_word_document(document_content, output_dir, timestamp, folder_name)
        else:
            # Fallback to markdown if docx not available
            doc_path = os.path.join(output_dir, f"Functional_Requirements_{folder_name}_{timestamp}.md")
            with open(doc_path, 'w', encoding='utf-8') as f:
                f.write(document_content)
            print("⚠️ Generated Markdown (install python-docx for Word format)")
        
        # Save summary
        self._save_summary(output_dir, timestamp)
        
        print(f"✅ Document generated: {doc_path}")
        return doc_path
    
    def _scan_biztalk_files(self, root_dir: str) -> Dict[str, Any]:
        """Smart file scanning with depth control"""
        print("🔍 Scanning BizTalk files...")
        
        analysis = {"orchestrations": [], "schemas": [], "maps": [], "bindings": [], "pipelines": []}
        
        if not os.path.exists(root_dir):
            return analysis
        
        # File patterns to search for
        patterns = {
            'orchestrations': '.odx',
            'schemas': '.xsd', 
            'maps': '.btm',
            'bindings': 'BindingInfo.xml',
            'pipelines': '.btp'
        }
        
        # Walk directory with depth control
        for root, dirs, files in os.walk(root_dir):
            current_depth = root.replace(root_dir, '').count(os.sep)
            if current_depth >= self.scan_depth:
                dirs.clear()  # Stop going deeper
                continue
            
            for filename in files:
                file_path = os.path.join(root, filename)
                
                # Check each pattern
                for category, pattern in patterns.items():
                    if self._matches_pattern(filename, pattern):
                        try:
                            info = self._analyze_file(file_path, category)
                            analysis[category].append(info)
                            self.analysis_results['processed_files'].append(file_path)
                        except Exception as e:
                            print(f"⚠️ Skipped {filename}: {e}")
        
        print(f"📊 Found {len(self.analysis_results['processed_files'])} files")
        return analysis
    
    def _matches_pattern(self, filename: str, pattern: str) -> bool:
        """Simple pattern matching"""
        return pattern in filename
    
    def _analyze_file(self, file_path: str, category: str) -> Dict[str, Any]:
        """Lightweight file analysis"""
        file_info = {
            "name": Path(file_path).stem,
            "path": file_path,
            "size": os.path.getsize(file_path),
            "category": category
        }
        
        # Try to extract basic XML info if possible
        if file_path.endswith(('.xml', '.xsd', '.btm')):
            try:
                tree = ET.parse(file_path)
                root = tree.getroot()
                file_info["elements_count"] = len(list(root.iter()))
                file_info["root_tag"] = root.tag.split('}')[-1] if '}' in root.tag else root.tag
            except:
                pass  # Skip XML parsing errors
        
        return file_info
    
    def _process_confluence(self, confluence_path: str) -> Dict[str, Any]:
        """Process Confluence extract"""
        print("📖 Processing Confluence extract...")
        
        try:
            if confluence_path.endswith('.json'):
                with open(confluence_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                return {"business_context": str(data)[:1000]}  # Limit size
            
            elif confluence_path.endswith(('.txt', '.md')):
                with open(confluence_path, 'r', encoding='utf-8') as f:
                    content = f.read()[:2000]  # Limit size
                
                # Use LLM to extract key info if available
                if self.llm:
                    return self._extract_with_llm(content)
                else:
                    return {"business_context": content}
        
        except Exception as e:
            print(f"❌ Confluence processing error: {e}")
        
        return {}
    
    def _extract_with_llm(self, content: str) -> Dict[str, Any]:
        """Extract structured info using LLM - 100% AI-powered"""
        if not self.llm:
            return {"business_context": content[:500]}  # Basic truncation only
        
        try:
            prompt = f"Extract key business requirements from this text in 3-5 bullet points:\n\n{content}"
            
            response = self.llm.chat.completions.create(
                model=self.config["groq_settings"]["model"],
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1,
                max_tokens=200
            )
            
            return {"business_requirements": [response.choices[0].message.content]}
        
        except Exception as e:
            print(f"❌ LLM extraction error: {e}")
            return {"business_context": content[:500]}  # Basic truncation only, no hardcoded content
    
    def _generate_with_llm(self, biztalk_analysis: Dict, confluence_analysis: Dict) -> str:
        """Generate documentation using LLM"""
        try:
            prompt = get_functional_documentation_prompt(
                biztalk_analysis, confluence_analysis, self.analysis_results
            )
            
            response = self.llm.chat.completions.create(
                model=self.config["groq_settings"]["model"],
                messages=[
                    {"role": "system", "content": "You are an expert technical writer creating migration documentation."},
                    {"role": "user", "content": prompt}
                ],
                temperature=self.config["groq_settings"]["temperature"]
            )
            
            return response.choices[0].message.content
            
        except Exception as e:
            print(f"❌ LLM generation failed: {e}")
            return self._create_template_doc(biztalk_analysis, confluence_analysis)
    
    def _create_word_document(self, content: str, output_dir: str, timestamp: str, folder_name: str) -> str:
        """Create Word document from markdown content"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available")
        
        doc = Document()
        
        # Add title with folder name
        title = doc.add_heading(f'{folder_name} - BizTalk to IBM ACE Migration', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add subtitle
        subtitle = doc.add_heading('Functional Requirements Document', level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add document info
        doc.add_heading('Document Information', level=1)
        info_table = doc.add_table(rows=3, cols=2)
        info_table.style = 'Table Grid'
        
        info_data = [
            ['Generated', datetime.now().strftime("%Y-%m-%d %H:%M:%S")],
            ['Files Processed', str(len(self.analysis_results['processed_files']))],
            ['Scan Depth', str(self.scan_depth)]
        ]
        
        for i, (key, value) in enumerate(info_data):
            info_table.cell(i, 0).text = key
            info_table.cell(i, 1).text = value
        
        # Process markdown content into Word format
        self._add_markdown_to_doc(doc, content)
        
        # Save document with folder name
        doc_path = os.path.join(output_dir, f"Functional_Requirements_{folder_name}_{timestamp}.docx")
        doc.save(doc_path)
        
        return doc_path
    
    def _add_markdown_to_doc(self, doc: Document, content: str):
        """Convert markdown content to Word document format"""
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Handle headers
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=4)
            
            # Handle bullet points
            elif line.startswith('- '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
            
            # Handle numbered lists
            elif line.startswith(('1. ', '2. ', '3. ', '4. ', '5. ', '6. ', '7. ', '8. ', '9. ')):
                p = doc.add_paragraph(line[3:], style='List Number')
            
            # Handle bold text (simplified)
            elif '**' in line:
                p = doc.add_paragraph()
                # Simple bold text handling
                parts = line.split('**')
                for i, part in enumerate(parts):
                    if i % 2 == 0:
                        p.add_run(part)
                    else:
                        p.add_run(part).bold = True
            
            # Regular paragraphs
            else:
                if not line.startswith(('|', '```', '---')):  # Skip tables and code blocks
                    doc.add_paragraph(line)
    
    def _create_template_doc(self, biztalk_analysis: Dict, confluence_analysis: Dict) -> str:
        """Simple template-based document"""
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # Calculate totals
        totals = {k: len(v) for k, v in biztalk_analysis.items()}
        
        # Confluence section
        confluence_section = ""
        if confluence_analysis.get('business_requirements'):
            confluence_section = f"\n## Business Context\n{confluence_analysis['business_requirements'][0]}\n"
        
        return f"""# BizTalk to IBM ACE Migration - Functional Requirements

## Document Information
- **Generated**: {timestamp}
- **Files Processed**: {len(self.analysis_results['processed_files'])}
- **Scan Depth**: {self.scan_depth} levels

{confluence_section}

## Current State Analysis

### Component Summary
- **Orchestrations**: {totals['orchestrations']} files
- **Schemas**: {totals['schemas']} files
- **Maps**: {totals['maps']} files
- **Bindings**: {totals['bindings']} files
- **Pipelines**: {totals['pipelines']} files

### Key Components
{self._format_components(biztalk_analysis)}

## Future State (IBM ACE)

The target solution will use:
- **Message Flows** for {totals['orchestrations']} business processes
- **ESQL Procedures** for {totals['maps']} data transformations  
- **Message Models** for {totals['schemas']} schema definitions
- **Connector Nodes** for {totals['bindings']} integrations

## Migration Requirements

### Functional Requirements
- Preserve all business logic and process flows
- Maintain data transformation accuracy
- Support existing integration patterns
- Implement equivalent error handling

### Technical Requirements
- Convert BizTalk orchestrations to ACE message flows
- Transform map logic into ESQL procedures
- Migrate schemas to ACE message models
- Configure connector nodes for integrations

## Success Criteria
- 100% functional equivalence
- Performance parity or better
- Successful end-to-end testing
- Production deployment readiness

---
*Generated by BizTalk to ACE Migration Tool*
"""
    
    def _format_components(self, analysis: Dict) -> str:
        """Format component list"""
        lines = []
        for category, items in analysis.items():
            if items:
                lines.append(f"\n#### {category.title()}")
                for item in items[:5]:  # Show max 5 per category
                    lines.append(f"- {item.get('name', 'Unknown')}")
                if len(items) > 5:
                    lines.append(f"- ... and {len(items) - 5} more")
        return '\n'.join(lines)
    
    def _save_summary(self, output_dir: str, timestamp: str):
        """Save generation summary"""
        summary = {
            "timestamp": timestamp,
            "scan_depth": self.scan_depth,
            "files_processed": len(self.analysis_results['processed_files']),
            "file_breakdown": {
                "orchestrations": len([f for f in self.analysis_results['processed_files'] if f.endswith('.odx')]),
                "schemas": len([f for f in self.analysis_results['processed_files'] if f.endswith('.xsd')]),
                "maps": len([f for f in self.analysis_results['processed_files'] if f.endswith('.btm')]),
                "bindings": len([f for f in self.analysis_results['processed_files'] if 'BindingInfo' in f]),
                "pipelines": len([f for f in self.analysis_results['processed_files'] if f.endswith('.btp')])
            }
        }
        
        summary_path = os.path.join(output_dir, f"Summary_{timestamp}.json")
        with open(summary_path, 'w', encoding='utf-8') as f:
            json.dump(summary, f, indent=2)

def main():
    """Command line interface"""
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python functional_document_generator.py <biztalk_dir> [confluence_file]")
        print("Requirements: GROQ_API_KEY environment variable")
        print("Optional: python-docx for Word format (installs as .md if not available)")
        return
    
    biztalk_dir = sys.argv[1]
    confluence_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    try:
        if not DOCX_AVAILABLE:
            print("⚠️ python-docx not available - will generate Markdown format")
            print("   Install with: pip install python-docx")
        
        generator = FunctionalDocumentGenerator()
        result = generator.generate_functional_document(biztalk_dir, confluence_file)
        
        file_format = "Word (.docx)" if result.endswith('.docx') else "Markdown (.md)"
        print(f"🎉 Complete! Generated {file_format}")
        print(f"📄 Document: {result}")
        print(f"📊 Files: {len(generator.analysis_results['processed_files'])}")
        
    except Exception as e:
        print(f"❌ Error: {e}")

if __name__ == "__main__":
    main()